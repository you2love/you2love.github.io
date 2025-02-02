from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import random
from docx.enum.text import WD_ALIGN_PARAGRAPH


# 生成1到10之间（含1和10）的一个随机整数
def main():
    lower, height = 0, 20
    row_len, col_len = 17, 2

    doc = Document()

    # 获取文档的第一个节（section）
    section = doc.sections[0]

    # 设置页面上下边距（去除上下空白）
    section.top_margin = Inches(0)  # 上边距设置为 0
    section.bottom_margin = Inches(0)  # 下边距设置为 0

    # 设置页面边距（去除左右空白）
    section.left_margin = Inches(0)  # 左边距设置为 0
    section.right_margin = Inches(0)  # 右边距设置为 0

    # section.page_width = Pt(595)  # A4 纸宽度（21cm 转换为磅）
    # section.page_height = Pt(842)  # A4 纸高度（29.7cm 转换为磅）

    # 设置页面方向（横向或纵向）
    # section.orientation = WD_ORIENT.LANDSCAPE  # 横向

    # 添加标题，并设置字体大小和加粗
    # title = doc.add_heading('20以内加减法练习题', level=0)
    # title_run = title.runs[0]
    # title_run.font.size = Pt(26)  # 设置字体大小为 24 磅
    # title_run.bold = True  # 设置字体加粗
    # title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=row_len, cols=col_len)
    table.style = "Table Grid"

    # 添加练习题
    question, answer = 0, 0
    exercises = [
        " {q} + (   ) = {a}",
        " (   ) + {q} = {a}",
        " {a} - (   ) = {q}",
    ]

    fmt_len = len(exercises)
    for r in range(row_len):
        for c in range(col_len):
            question = random.randint(lower, height)
            answer = random.randint(lower, height)
            if question > answer:
                answer, question = question, answer

            fmt_pos = random.randint(0, fmt_len - 1)
            exercise = exercises[fmt_pos]

            table.cell(r, c).text = exercise.format(q=question, a=answer)
            # print('q', question, 'a', answer, 'row', table.cell(r,c).text )

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True  # 加粗
                    run.font.size = Pt(24)  # 设置字体大小为 14 磅
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
    # return
    doc.save("数学题.docx")


if __name__ == "__main__":
    main()
